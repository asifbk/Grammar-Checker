import tkinter as tk
from tkinter import scrolledtext, filedialog, messagebox, ttk
from threading import Thread
import difflib
import pyperclip
from docx import Document
from ollama import chat
import re
from PIL import Image, ImageTk
import platform
import tempfile
import os
import subprocess
import time
from datetime import datetime
import numpy as np
from collections import Counter

# ---------------- Globals ---------------- #
is_reading = False
tts_button = None
start_time = None

# Available models - only the three specified
AVAILABLE_MODELS = [
    "llama3.2", 
    "mistral",
    "deepseek-r1:7b"
]

# ---------------- Evaluation Metrics Functions ---------------- #
def calculate_gleu(reference, hypothesis, n=4):
    """
    Calculate GLEU score between reference and hypothesis texts.
    GLEU is based on BLEU but designed for grammatical error correction.
    """
    def get_ngrams(text, n):
        words = text.split()
        ngrams = []
        for i in range(len(words) - n + 1):
            ngrams.append(tuple(words[i:i + n]))
        return Counter(ngrams)
    
    ref_ngrams = get_ngrams(reference, n)
    hyp_ngrams = get_ngrams(hypothesis, n)
    
    # Count matching n-grams
    match_count = 0
    total_count = 0
    
    for ngram, count in hyp_ngrams.items():
        total_count += count
        if ngram in ref_ngrams:
            match_count += min(count, ref_ngrams[ngram])
    
    if total_count == 0:
        return 0.0
    
    return match_count / total_count

def calculate_errant_precision(original, corrected, reference):
    """
    Calculate ERRANT-style precision for grammatical error correction.
    Measures how many corrections were actually needed.
    """
    # Simple implementation - count true corrections
    original_tokens = original.lower().split()
    corrected_tokens = corrected.lower().split()
    reference_tokens = reference.lower().split()
    
    # Count changes between original and corrected
    changes = 0
    true_corrections = 0
    
    # Simple diff-based approach
    orig_words = original.lower().split()
    corr_words = corrected.lower().split()
    ref_words = reference.lower().split()
    
    # Count words that changed and check if they match reference
    for i in range(min(len(orig_words), len(corr_words))):
        if i < len(ref_words):
            if orig_words[i] != corr_words[i]:
                changes += 1
                if corr_words[i] == ref_words[i]:
                    true_corrections += 1
    
    if changes == 0:
        return 1.0  # No changes made, but if original was correct, precision is high
    
    return true_corrections / changes if changes > 0 else 1.0

def calculate_grade_score(original, corrected):
    """
    Calculate a composite grade score based on multiple metrics.
    Higher score means better grammatical quality.
    """
    score = 0.0
    
    # 1. Sentence structure improvement
    orig_sentences = re.split(r'[.!?]+', original)
    corr_sentences = re.split(r'[.!?]+', corrected)
    
    if len(corr_sentences) >= len(orig_sentences):
        score += 0.2
    
    # 2. Word count ratio (penalize excessive shortening/lengthening)
    orig_words = len(original.split())
    corr_words = len(corrected.split())
    word_ratio = min(corr_words, orig_words) / max(corr_words, orig_words) if max(corr_words, orig_words) > 0 else 1.0
    score += word_ratio * 0.2
    
    # 3. Punctuation improvement
    orig_punct = len(re.findall(r'[,.!?;:]', original))
    corr_punct = len(re.findall(r'[,.!?;:]', corrected))
    if corr_punct >= orig_punct:
        score += 0.2
    
    # 4. Capitalization improvement
    orig_caps = len(re.findall(r'\b[A-Z][a-z]+\b', original))
    corr_caps = len(re.findall(r'\b[A-Z][a-z]+\b', corrected))
    if corr_caps >= orig_caps:
        score += 0.2
    
    # 5. Basic grammar checks
    grammar_issues_orig = count_grammar_issues(original)
    grammar_issues_corr = count_grammar_issues(corrected)
    if grammar_issues_corr < grammar_issues_orig:
        score += 0.2
    
    return min(score, 1.0)  # Normalize to 0-1

def count_grammar_issues(text):
    """Count basic grammar issues in text"""
    issues = 0
    
    # Check for common errors
    if re.search(r'\bi\s+', text.lower()):  # lowercase 'i' as subject
        issues += 1
    if re.search(r'\b(a|an|the)\s+[aeiouAEIOU]', text):  # article issues
        issues += 1
    if re.search(r'\b(he|she|it)\s+(go|have|do)\b', text.lower()):  # verb agreement
        issues += 1
    if re.search(r'\b(we|they|you)\s+(goes|has|does)\b', text.lower()):  # verb agreement
        issues += 1
    if re.search(r'\b(much|many)\s+\w+s\b', text.lower()):  # much/many issues
        issues += 1
    
    return issues

def calculate_comprehensive_score(original, corrected, reference=None):
    """
    Calculate comprehensive evaluation metrics.
    If reference is provided, use it for GLEU and ERRANT.
    Otherwise, use original as reference for basic metrics.
    """
    if reference is None:
        reference = original  # For basic evaluation without gold standard
    
    # Calculate all metrics
    gleu_score = calculate_gleu(reference, corrected)
    errant_precision = calculate_errant_precision(original, corrected, reference)
    grade_score = calculate_grade_score(original, corrected)
    word_count = len(corrected.split())
    
    # Composite score (weighted average)
    composite_score = (gleu_score * 0.3 + errant_precision * 0.4 + grade_score * 0.3) * 100
    
    return {
        'gleu': gleu_score,
        'errant_precision': errant_precision,
        'grade_score': grade_score,
        'composite_score': composite_score,
        'word_count': word_count
    }

# ---------------- Improved Model Management Functions ---------------- #
def get_installed_models():
    """Get list of models currently installed in Ollama"""
    installed_models = []
    try:
        # Method 1: Try ollama list command
        result = subprocess.run(["ollama", "list"], capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            lines = result.stdout.strip().split('\n')
            for line in lines[1:]:  # Skip header
                if line.strip():
                    parts = line.split()
                    if parts:
                        model_name = parts[0]
                        installed_models.append(model_name)
        
        # Method 2: Try direct API call to Ollama for each model
        test_models = ["llama3.2", "mistral", "deepseek-r1", "deepseek-r1:7b", "deepseek-r1:14b"]
        for model in test_models:
            try:
                response = chat(model=model, messages=[{"role": "user", "content": "hi"}], options={"timeout": 3})
                if model not in installed_models:
                    installed_models.append(model)
            except:
                pass
        
        print(f"Detected installed models: {installed_models}")
        return installed_models
        
    except Exception as e:
        print(f"Error detecting models: {e}")
        return []

def install_model(model_name):
    """Install a model using Ollama"""
    try:
        status_label.config(text=f"Downloading {model_name}...")
        result = subprocess.run(["ollama", "pull", model_name], capture_output=True, text=True, timeout=300)
        if result.returncode == 0:
            messagebox.showinfo("Success", f"Model {model_name} installed successfully!")
            refresh_model_list()
            return True
        else:
            messagebox.showerror("Error", f"Failed to install {model_name}: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        messagebox.showerror("Error", f"Download timeout for {model_name}")
        return False
    except Exception as e:
        messagebox.showerror("Error", f"Failed to install {model_name}: {str(e)}")
        return False
    finally:
        status_label.config(text="")

def refresh_model_list():
    """Refresh the model dropdown with installed models"""
    try:
        installed_models = get_installed_models()
        
        # Update dropdown menu
        menu = model_menu["menu"]
        menu.delete(0, "end")
        
        # Check which of our available models are installed
        found_installed = False
        for model in AVAILABLE_MODELS:
            # Check if model is installed (exact match or contains)
            is_installed = any(installed.startswith(model) or model in installed for installed in installed_models)
            
            if is_installed:
                menu.add_command(label=model, command=lambda value=model: model_var.set(value))
                found_installed = True
                print(f"Found installed: {model}")
        
        # If no installed models found, show all as available
        if not found_installed:
            for model in AVAILABLE_MODELS:
                menu.add_command(label=f"{model} (Not installed)", 
                               command=lambda value=model: install_model_dialog(value))
        else:
            # Add separator and show unavailable models
            menu.add_separator()
            for model in AVAILABLE_MODELS:
                is_installed = any(installed.startswith(model) or model in installed for installed in installed_models)
                if not is_installed:
                    menu.add_command(label=f"{model} (Not installed)", 
                                   command=lambda value=model: install_model_dialog(value))
        
        # Update the current selection if it's installed
        current_model = model_var.get()
        if "(Not installed)" in current_model:
            model_name = current_model.replace(" (Not installed)", "")
            if any(installed.startswith(model_name) or model_name in installed for installed in installed_models):
                model_var.set(model_name)
                
    except Exception as e:
        print(f"Error refreshing model list: {e}")
        # Fallback: show all models as available
        menu = model_menu["menu"]
        menu.delete(0, "end")
        for model in AVAILABLE_MODELS:
            menu.add_command(label=model, command=lambda value=model: model_var.set(value))

def install_model_dialog(model_name):
    """Ask user if they want to install a missing model"""
    response = messagebox.askyesno(
        "Model Not Found", 
        f"Model '{model_name}' is not installed.\n\nDo you want to download it now? (This may take several minutes)"
    )
    if response:
        Thread(target=lambda: install_model(model_name), daemon=True).start()

# ---------------- Robust TTS Implementation ---------------- #
class RobustTTS:
    def __init__(self):
        self.is_reading = False
        self.engine_type = None
        self.engine = None
        self.initialize_tts()
    
    def initialize_tts(self):
        """Initialize TTS with multiple fallback options"""
        try:
            import pyttsx3
            self.engine = pyttsx3.init()
            voices = self.engine.getProperty('voices')
            if voices:
                self.engine.setProperty('voice', voices[0].id)
            self.engine.setProperty('rate', 150)
            self.engine_type = "pyttsx3"
            return True
        except Exception as e:
            print(f"pyttsx3 failed: {e}")
        
        system = platform.system()
        try:
            if system == "Windows":
                import win32com.client
                self.engine = win32com.client.Dispatch("SAPI.SpVoice")
                self.engine_type = "windows"
                return True
            elif system == "Darwin":
                self.engine_type = "macos"
                return True
            elif system == "Linux":
                self.engine_type = "linux"
                return True
        except Exception as e:
            print(f"System TTS failed: {e}")
        
        try:
            from gtts import gTTS
            import pygame
            pygame.mixer.init()
            self.engine_type = "gtts"
            return True
        except ImportError:
            pass
        
        return False
    
    def speak(self, text):
        if not self.engine_type:
            return False
        try:
            if self.engine_type == "pyttsx3":
                self.engine.say(text)
                self.engine.runAndWait()
            elif self.engine_type == "windows":
                self.engine.Speak(text)
            elif self.engine_type == "macos":
                import subprocess
                subprocess.run(["say", text], capture_output=True)
            elif self.engine_type == "linux":
                import subprocess
                try:
                    subprocess.run(["espeak", text], capture_output=True)
                except FileNotFoundError:
                    subprocess.run(["spd-say", text], capture_output=True)
            elif self.engine_type == "gtts":
                self._speak_gtts(text)
            return True
        except Exception as e:
            print(f"TTS speak error: {e}")
            return False
    
    def _speak_gtts(self, text):
        try:
            from gtts import gTTS
            import pygame
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as tmp_file:
                temp_filename = tmp_file.name
            
            chunks = self._split_text(text, 200)
            for chunk in chunks:
                tts = gTTS(text=chunk, lang='en', slow=False)
                tts.save(temp_filename)
                pygame.mixer.music.load(temp_filename)
                pygame.mixer.music.play()
                while pygame.mixer.music.get_busy():
                    pygame.time.wait(100)
            
            pygame.mixer.music.unload()
            os.unlink(temp_filename)
        except Exception as e:
            print(f"Google TTS error: {e}")
    
    def _split_text(self, text, max_length):
        words = text.split()
        chunks = []
        current_chunk = []
        for word in words:
            if len(' '.join(current_chunk + [word])) <= max_length:
                current_chunk.append(word)
            else:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [word]
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        return chunks
    
    def stop(self):
        try:
            if self.engine_type == "pyttsx3" and self.engine:
                self.engine.stop()
            elif self.engine_type == "gtts":
                import pygame
                pygame.mixer.music.stop()
        except Exception as e:
            print(f"TTS stop error: {e}")

# Global TTS engine
tts_engine = RobustTTS()

# ---------------- Grammar Checker Logic ---------------- #
def run_correction(input_text, language="English", style="Default"):
    """Corrects grammar and spelling using the Ollama model."""
    prompt = (
        f"Correct the grammar and spelling of the following {language} text. "
        f"Return only the corrected text in a {style.lower()} style.\n\nText: \"{input_text}\""
    )
    try:
        selected_model = model_var.get()
        # Remove "(Not installed)" from model name if present
        selected_model = selected_model.split(" (Not installed)")[0]
        
        # Try different DeepSeek variations if the main one fails
        if selected_model == "deepseek-r1:7b":
            try:
                response = chat(
                    model=selected_model,
                    messages=[{"role": "user", "content": prompt}],
                    options={"timeout": 45}
                )
                return response['message']['content'].strip()
            except:
                # Try alternative DeepSeek names
                for alt_model in ["deepseek-r1", "deepseek-r1:14b", "deepseek-coder"]:
                    try:
                        response = chat(
                            model=alt_model,
                            messages=[{"role": "user", "content": prompt}],
                            options={"timeout": 45}
                        )
                        return response['message']['content'].strip()
                    except:
                        continue
                raise Exception("DeepSeek model not working. Try 'ollama pull deepseek-r1:7b'")
        else:
            response = chat(
                model=selected_model,
                messages=[{"role": "user", "content": prompt}],
                options={"timeout": 45}
            )
            return response['message']['content'].strip()
            
    except Exception as e:
        if "not found" in str(e).lower():
            raise Exception(f"Model '{model_var.get()}' not found. Please install it first.")
        else:
            raise Exception(f"Model error: {str(e)}")

# ---------------- Stats Functions ---------------- #
def count_words(text):
    words = re.findall(r"\w+", text)
    return len(words)

def flesch_kincaid(text):
    sentences = max(1, text.count('.') + text.count('!') + text.count('?'))
    words = count_words(text)
    syllables = sum(count_syllables(word) for word in re.findall(r"\w+", text))
    if words == 0:
        return 0
    return round(0.39 * (words / sentences) + 11.8 * (syllables / words) - 15.59, 2)

def count_syllables(word):
    word = word.lower()
    vowels = "aeiouy"
    count = 0
    prev_char_was_vowel = False
    for char in word:
        if char in vowels:
            if not prev_char_was_vowel:
                count += 1
            prev_char_was_vowel = True
        else:
            prev_char_was_vowel = False
    if word.endswith("e"):
        count = max(1, count - 1)
    return max(1, count)

# ---------------- GUI Functions ---------------- #
def correct_grammar_threaded():
    def task():
        global start_time
        start_time = time.time()
        
        input_text = input_box.get("1.0", tk.END).strip()
        if not input_text:
            messagebox.showwarning("Warning", "Please enter some text to correct.")
            return
        
        root.after(0, lambda: status_label.config(text="Processing..."))
        root.after(0, progress.start)
        
        try:
            lang = language_var.get()
            style = style_var.get()
            corrected_text = run_correction(input_text, language=lang, style=style)
            
            # Calculate completion time
            end_time = time.time()
            completion_time = end_time - start_time
            
            # Calculate comprehensive evaluation metrics
            metrics = calculate_comprehensive_score(input_text, corrected_text)
            
            # Update GUI in main thread
            root.after(0, update_output, corrected_text, input_text, completion_time, metrics)
            
        except Exception as e:
            end_time = time.time()
            completion_time = end_time - start_time if start_time else 0
            root.after(0, show_error, f"Correction failed: {str(e)}", completion_time)
        finally:
            root.after(0, progress.stop)

    Thread(target=task, daemon=True).start()

def update_output(corrected_text, original_text, completion_time, metrics):
    output_box.delete("1.0", tk.END)
    output_box.insert(tk.END, corrected_text)
    highlight_corrections(original_text, corrected_text)
    
    words = count_words(corrected_text)
    grade = flesch_kincaid(corrected_text)
    
    # Update stats with all metrics
    stats_text = (
        f"Words: {metrics['word_count']} | "
        f"Time: {completion_time:.2f}s | "
        f"GLEU: {metrics['gleu']:.3f} | "
        f"ERRANT: {metrics['errant_precision']:.3f} | "
        f"Grade: {metrics['grade_score']:.3f} | "
        f"Score: {metrics['composite_score']:.1f}%"
    )
    stats_label.config(text=stats_text)
    status_label.config(text="")
    
    # Show completion message with time
    current_time = datetime.now().strftime("%H:%M:%S")
    completion_label.config(text=f"Completed at {current_time} | Processing time: {completion_time:.2f} seconds")

def show_error(message, completion_time=0):
    messagebox.showerror("Error", message)
    status_label.config(text="")
    if completion_time > 0:
        completion_label.config(text=f"Failed after {completion_time:.2f} seconds")
    else:
        completion_label.config(text="")

def highlight_corrections(original, corrected):
    output_box.tag_delete("correction")
    output_box.tag_config("correction", foreground="red")
    original_words = original.split()
    corrected_words = corrected.split()
    s = difflib.SequenceMatcher(None, original_words, corrected_words)
    current_char_index = 0
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        corrected_chunk = " ".join(corrected_words[j1:j2])
        if tag == 'replace' or tag == 'delete' or tag == 'insert':
            end_char_index = current_char_index + len(corrected_chunk)
            output_box.tag_add("correction", f"1.0+{current_char_index}c", f"1.0+{end_char_index}c")
        current_char_index += len(corrected_chunk)
        if corrected_chunk and current_char_index < len(corrected):
             current_char_index += 1

def clear_text():
    input_box.delete("1.0", tk.END)
    output_box.delete("1.0", tk.END)
    stats_label.config(text="Words: 0 | Time: 0.00s | GLEU: 0.000 | ERRANT: 0.000 | Grade: 0.000 | Score: 0.0%")
    status_label.config(text="")
    completion_label.config(text="")

def copy_output():
    corrected_text = output_box.get("1.0", tk.END).strip()
    if not corrected_text:
        messagebox.showwarning("Warning", "No text to copy!")
        return
    pyperclip.copy(corrected_text)
    messagebox.showinfo("Copied", "Corrected text copied to clipboard!")

def save_output():
    corrected = output_box.get("1.0", tk.END).strip()
    if not corrected:
        messagebox.showwarning("Warning", "No text to save!")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")])
    if not file_path:
        return
    try:
        if file_path.endswith(".txt"):
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(corrected)
        elif file_path.endswith(".docx"):
            doc = Document()
            doc.add_paragraph(corrected)
            doc.save(file_path)
        messagebox.showinfo("Saved", "File saved successfully!")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save file: {str(e)}")

def swap_text():
    input_text = input_box.get("1.0", tk.END).strip()
    output_text = output_box.get("1.0", tk.END).strip()
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, output_text)
    output_box.delete("1.0", tk.END)
    output_box.insert(tk.END, input_text)

def load_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("All files", "*.*")]
    )
    if file_path:
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            input_box.delete("1.0", tk.END)
            input_box.insert(tk.END, text)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load file: {str(e)}")

# ---------------- TTS Functions ---------------- #
def toggle_read_aloud():
    global is_reading, tts_button
    corrected = output_box.get("1.0", tk.END).strip()
    if not corrected:
        messagebox.showwarning("Warning", "No text to read!")
        return
    if is_reading:
        pause_reading()
        tts_button.config(text="▶ Play", fg="blue")
    else:
        read_aloud(corrected)
        tts_button.config(text="⏸ Pause", fg="yellow")

def read_aloud(text):
    global is_reading
    if not is_reading and tts_engine.engine_type:
        is_reading = True
        Thread(target=lambda: tts_thread(text), daemon=True).start()
    elif not tts_engine.engine_type:
        messagebox.showwarning("TTS Unavailable", "Text-to-speech is not available on this system.")

def tts_thread(text):
    global is_reading, tts_button
    try:
        success = tts_engine.speak(text)
        if not success:
            root.after(0, lambda: messagebox.showwarning("TTS Error", "Text-to-speech failed to play."))
    except Exception as e:
        root.after(0, lambda: messagebox.showerror("TTS Error", f"Text-to-speech failed: {e}"))
    finally:
        is_reading = False
        root.after(0, lambda: tts_button.config(text="▶ Play", fg="blue"))

def pause_reading():
    global is_reading
    tts_engine.stop()
    is_reading = False

def stop_reading():
    global is_reading, tts_button
    tts_engine.stop()
    is_reading = False
    tts_button.config(text="▶ Play", fg="blue")

# ---------------- GUI Setup ---------------- #
root = tk.Tk()
root.title("UALR Advanced Grammar Checker")
root.geometry("1200x750")  # Increased width for more metrics

# Frames
top_frame = tk.Frame(root)
top_frame.pack(side=tk.TOP, fill=tk.X, pady=5)

controls_frame = tk.Frame(root)
controls_frame.pack(side=tk.TOP, fill=tk.X, pady=5)

text_frame = tk.Frame(root)
text_frame.pack(fill=tk.BOTH, expand=True)

bottom_frame = tk.Frame(root)
bottom_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=5)

# Model Selection with only the three specified models
model_var = tk.StringVar(value="llama3.2")
tk.Label(controls_frame, text="AI Model:").pack(side=tk.LEFT, padx=5)
model_menu = tk.OptionMenu(controls_frame, model_var, "Loading...")
model_menu.pack(side=tk.LEFT, padx=5)

# Refresh models button
tk.Button(controls_frame, text="🔄 Refresh Models", command=refresh_model_list).pack(side=tk.LEFT, padx=5)

# Manual override button - in case detection fails
def force_models_installed():
    """Force set all models as installed"""
    for model in AVAILABLE_MODELS:
        model_var.set(model)
        break  # Set to first model
    messagebox.showinfo("Manual Override", "Models set as installed. You can now try using them.")

tk.Button(controls_frame, text="🔧 Manual Override", command=force_models_installed).pack(side=tk.LEFT, padx=5)

# Language & Style dropdown
language_var = tk.StringVar(value="English")
style_var = tk.StringVar(value="Default")
tk.Label(controls_frame, text="Language:").pack(side=tk.LEFT, padx=5)
tk.OptionMenu(controls_frame, language_var, "English", "Spanish", "French", "German", "Italian", "Portuguese", "Dutch", "Russian", "Japanese", "Chinese", "Korean", "Arabic", "Hindi", "Bengali").pack(side=tk.LEFT)
tk.Label(controls_frame, text="Style:").pack(side=tk.LEFT, padx=5)
tk.OptionMenu(controls_frame, style_var, "Default", "Formal", "Academic", "Casual", "Professional").pack(side=tk.LEFT)

# Input (Left)
input_label = tk.Label(text_frame, text="Input Text:")
input_label.grid(row=0, column=0, sticky="w")
input_box = scrolledtext.ScrolledText(text_frame, width=60, height=20)
input_box.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

# Output (Right)
output_label = tk.Label(text_frame, text="Corrected Text:")
output_label.grid(row=0, column=1, sticky="w")
output_box = scrolledtext.ScrolledText(text_frame, width=60, height=20)
output_box.grid(row=1, column=1, padx=5, pady=5, sticky="nsew")

# Allow resizing
text_frame.columnconfigure(0, weight=1)
text_frame.columnconfigure(1, weight=1)

# Buttons
tk.Button(top_frame, text="Correct Grammar", command=correct_grammar_threaded).pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="Clear", command=clear_text).pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="Swap Text", command=swap_text).pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="Copy", command=copy_output).pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="Save", command=save_output).pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="Load File", command=load_file).pack(side=tk.LEFT, padx=5)

# TTS Buttons
tts_button = tk.Button(top_frame, text="▶ Play", font=("Arial", 12), fg="blue", command=toggle_read_aloud)
tts_button.pack(side=tk.LEFT, padx=5)
tk.Button(top_frame, text="⏹ Stop", font=("Arial", 12), fg="red", command=stop_reading).pack(side=tk.LEFT, padx=5)

# Progress bar
progress = ttk.Progressbar(bottom_frame, mode='indeterminate')
progress.pack(anchor="w", fill=tk.X, padx=5, pady=2)

# Status + Comprehensive Metrics
status_label = tk.Label(bottom_frame, text="", fg="blue")
status_label.pack(anchor="w")

# Comprehensive stats label
stats_label = tk.Label(bottom_frame, 
                      text="Words: 0 | Time: 0.00s | GLEU: 0.000 | ERRANT: 0.000 | Grade: 0.000 | Score: 0.0%",
                      fg="darkgreen",
                      font=("Arial", 10, "bold"))
stats_label.pack(anchor="w")

# Completion time label
completion_label = tk.Label(bottom_frame, text="", fg="purple")
completion_label.pack(anchor="w")

# UALR Logo (bottom right)
try:
    logo_img = Image.open("ualr_logo1.png")
    logo_img = logo_img.resize((300, 100), Image.Resampling.LANCZOS)
    logo_photo = ImageTk.PhotoImage(logo_img)
    logo_label = tk.Label(bottom_frame, image=logo_photo)
    logo_label.image = logo_photo
    logo_label.pack(side=tk.RIGHT, padx=10, anchor="e")
except Exception as e:
    print("Logo not loaded:", e)

# TTS status
tts_status = "Available" if tts_engine.engine_type else "Unavailable"
tts_status_label = tk.Label(bottom_frame, text=f"TTS: {tts_status}", fg="green" if tts_engine.engine_type else "red")
tts_status_label.pack(anchor="w")

# Initialize model list on startup
root.after(1000, refresh_model_list)  # Delay to ensure GUI is loaded

root.mainloop()