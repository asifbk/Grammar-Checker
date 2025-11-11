import tkinter as tk
from tkinter import scrolledtext, filedialog, messagebox, ttk
from threading import Thread
import difflib
import webbrowser
import pyperclip
from docx import Document
try:
    from ollama import chat
except Exception:
    chat = None
import re
from PIL import Image, ImageTk
import platform
import tempfile
import os
import subprocess
import time
from datetime import datetime
import numpy as np
from collections import Counter
import csv

# ---------------- Globals ---------------- #
start_time = None
data_log = []

# Added a composite score column for logging
LOG_FIELDNAMES = [
    'timestamp', 'words', 'time', 'f05', 'err', 'gleu',
    'grade_level', 'composite_score', 'language', 'style', 'model'
]

AVAILABLE_MODELS = ["llama3.2", "mistral", "deepseek-r1:7b", "gemma3"]


# --- Metrics Functions ---
def count_syllables(w):
    w = w.lower().strip()
    if not w: return 0
    # Simple syllable counting logic (approximation)
    vowels = "aeiouy"
    count = 0
    prev_char_is_vowel = False

    for char in w:
        if char in vowels:
            if not prev_char_is_vowel:
                count += 1
            prev_char_is_vowel = True
        else:
            prev_char_is_vowel = False

    # Handle silent 'e' at the end of a word
    if w.endswith('e') and len(w) > 1 and w[-2] not in vowels:
        count -= 1

    return max(1, count)

def count_words(t):
    return len(re.findall(r'\w+', t))

def calculate_gleu(reference, hypothesis, n=4):
    def get_ngrams(text, n):
        words = text.split()
        ngrams = []
        for i in range(len(words) - n + 1):
            ngrams.append(tuple(words[i:i + n]))
        return Counter(ngrams)

    # Use a simpler reference for GLEU in a grammar checker context: the hypothesis itself
    # This is a common practice when a perfect "reference" corrected text isn't available.
    # We compare the n-grams in the hypothesis to the n-grams in the original text (as a loose reference)
    # or, more correctly for GLEU, you need multiple human-corrected references.
    # Since we don't have multiple references, we'll use a simplified metric against the original.
    # However, for a true GLEU as intended by the metric, a single hypothesis is compared against *multiple* references.
    # To keep the code running with the existing logic, we compare against the original (as the "reference").
    # This will heavily penalize any changes, which is generally incorrect for GEC metrics, but is required for the code's structure.
    # For a more practical approximation, let's use the standard definition for a single reference comparison.
    ref_ngrams = get_ngrams(reference, n)
    hyp_ngrams = get_ngrams(hypothesis, n)
    
    match = sum(min(hyp_ngrams[g], ref_ngrams.get(g, 0)) for g in hyp_ngrams)
    total = sum(hyp_ngrams.values())
    return match / total if total else 0.0

def calculate_err(original, corrected):
    sm = difflib.SequenceMatcher(None, original.split(), corrected.split())
    # Count of edit operations (substitutions, insertions, deletions)
    edits = sum(max(i2-i1, j2-j1) for tag,i1,i2,j1,j2 in sm.get_opcodes() if tag!='equal')
    return (edits / max(1,len(original.split()))) * 100

def calculate_precision(original, corrected):
    # Precision: Corrected text words matching the original text's meaning/intent (Approximation)
    # Higher precision means fewer unnecessary changes.
    sm = difflib.SequenceMatcher(None, original.split(), corrected.split())
    equal_words = sum(i2-i1 for tag,i1,i2,j1,j2 in sm.get_opcodes() if tag=='equal')
    total_corrected_words = len(corrected.split())
    return equal_words / max(1, total_corrected_words)

def calculate_recall(original, corrected):
    # Recall: Original text words needing correction that were corrected (Approximation)
    # Higher recall means more errors were caught and fixed.
    sm = difflib.SequenceMatcher(None, original.split(), corrected.split())
    
    # Simple approximation: errors caught = total original words - words remaining unchanged
    # A proper GEC metric requires a 'gold standard' reference to count true positives, false negatives, etc.
    # For this simplified model, we'll use the number of edits as a proxy for 'corrections made'.
    edits = sum(max(i2-i1, j2-j1) for tag,i1,i2,j1,j2 in sm.get_opcodes() if tag!='equal')
    
    # A very rough approximation of potential errors in original text:
    # Let's assume potential errors is proportional to the number of words.
    total_original_words = len(original.split())
    
    # A better proxy for Recall in this context is just the proportion of the original text that was changed.
    # This is often used when a gold reference is missing.
    return edits / max(1, total_original_words)


def calculate_f05(p, r):
    b2 = 0.25  # beta squared for F0.5
    return (1+b2)*p*r/((b2*p)+r) if p+r>0 else 0.0

def calculate_grade_level(text):
    """Calculate Flesch-Kincaid Grade Level for the given text"""
    # Count total words
    words = re.findall(r'\w+', text)
    total_words = len(words)
    
    # Count total sentences (ending with . ! ?)
    sentences = re.split(r'[.!?]+', text)
    total_sentences = len([s for s in sentences if s.strip()])
    
    # Count total syllables
    total_syllables = sum(count_syllables(word) for word in words)
    
    # Avoid division by zero
    if total_words == 0 or total_sentences == 0:
        return 0.0
    
    # Flesch-Kincaid Grade Level formula
    # Note: The UI description is confusing. This function implements Flesch-Kincaid, 
    # not the Flesch Reading Ease score mentioned in the UI description's formula.
    # The actual F-K Grade Level formula is: 
    grade_level = 0.39 * (total_words / total_sentences) + 11.8 * (total_syllables / total_words) - 15.59
    
    return max(0, round(grade_level, 2))

def calculate_metrics(original, corrected, processing_time):
    """Calculate all GEC and readability metrics"""
    
    # 1. Word Count (of original text)
    word_count = count_words(original)

    # 2. GLEU
    gleu = calculate_gleu(original, corrected)

    # 3. ERR
    err = calculate_err(original, corrected)

    # 4. F0.5 (requires Precision and Recall)
    p = calculate_precision(original, corrected)
    r = calculate_recall(original, corrected)
    f05 = calculate_f05(p, r)

    # 5. Grade Level (Flesch-Kincaid)
    grade_level = calculate_grade_level(corrected)
    
    # 6. Composite Score (Simple weighted average, higher is better)
    # Weights are arbitrary for demonstration:
    # F0.5 (high weight on precision)
    # GLEU (high weight on fluency/consistency)
    # (100 - ERR) (lower ERR is better, so maximize 100-ERR)
    # Time is not included in the composite score calculation, only for logging.
    
    # Normalize ERR (lower is better, so higher 100-ERR is better)
    normalized_err = max(0.0, 100.0 - err)
    
    # Simple linear combination for a score from 0 to 100
    # Weights: F0.5 (40%), GLEU (40%), 100-ERR (20%)
    composite_score = (f05 * 0.4 + gleu * 0.4 + (normalized_err / 100.0) * 0.2) * 100
    
    metrics = {
        'word_count': word_count,
        'f05': f05,
        'err': err,
        'gleu': gleu,
        'grade_level': grade_level,
        'composite_score': composite_score
    }
    return metrics

def get_installed_models():
    try:
        r=subprocess.run(["ollama","list"],capture_output=True,text=True,timeout=8)
        if r.returncode==0:
            return [l.split()[0] for l in r.stdout.strip().split('\n')[1:] if l.strip()]
    except Exception:pass
    return []

def refresh_model_list():
    def task():
        try:
            installed=get_installed_models()
            menu=model_menu["menu"];menu.delete(0,"end")
            for m in AVAILABLE_MODELS:
                lbl=m if any(m in i for i in installed) else f"{m} (Not installed)"
                menu.add_command(label=lbl,command=lambda v=lbl:model_var.set(v))
            root.after(0,lambda:status_label.config(text="Model list refreshed."))
        except Exception as e:
            root.after(0,lambda:status_label.config(text="Failed to refresh models."))
    Thread(target=task,daemon=True).start()

def run_correction(t, language="English", style="Default"):
    if chat is None: raise Exception("Ollama client not available.")
    
    # REFINED PROMPT: Increased emphasis on 'ONLY' and removed empty line after colon.
    prompt = f"""You are a strict, non-conversational grammar and spelling corrector. 
You MUST return ONLY the corrected text. DO NOT add any commentary, introductory phrases (like "Here is the corrected version"), labels (like "Corrected text:"), or any other text.
Correct ONLY the grammar, spelling, and punctuation of the following {language} text.
Maintain the {style.lower()} style.

Original text: "{t}"

Corrected text:""" # The prompt ends immediately after the colon
    
    m = model_var.get().split(" (Not installed)")[0]
    r = chat(model=m, messages=[{"role": "user", "content": prompt}], options={"timeout": 120})
    
    corrected_text = r['message']['content'].strip()
    
    # POST-PROCESSING: Aggressively clean up potential prefixes/suffixes the model might still add
    
    # 1. Remove common prefixes that still might sneak in
    prefixes_to_remove = [
        "Corrected text:", 
        "Sentence:", 
        "The corrected text is:", 
        "Here is the corrected text:"
    ]
    for prefix in prefixes_to_remove:
        if corrected_text.startswith(prefix):
            corrected_text = corrected_text[len(prefix):].strip()
            
    # 2. Strip surrounding quotes if the model wrapped the output
    if corrected_text.startswith('"') and corrected_text.endswith('"'):
        corrected_text = corrected_text[1:-1].strip()

    return corrected_text.strip()
# --------- Optimized Highlight Function ---------
def highlight_corrections(original, corrected):
    """Efficiently highlight only changed words in long texts."""
    output_box.tag_delete("correction")
    output_box.tag_config("correction", foreground="red")

    # Use SequenceMatcher to find differences at the character level
    sm = difflib.SequenceMatcher(None, original, corrected)
    
    # Track which parts of the corrected text are part of a change
    change_spans = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            continue
        # j1 to j2 is the span in the corrected text that is different
        change_spans.append((j1, j2))

    # Iterate over words in the corrected text
    for match in re.finditer(r'\S+', corrected):
        word_start, word_end = match.span()
        
        # Check if the current word overlaps with any detected change span
        is_changed = False
        for cs, ce in change_spans:
            # Overlap exists if the max of the starts is less than the min of the ends
            if max(word_start, cs) < min(word_end, ce):
                is_changed = True
                break
        
        if is_changed:
            output_box.tag_add("correction", f"1.0+{word_start}c", f"1.0+{word_end}c")

    # Use character-level SequenceMatcher to find differences
    s = difflib.SequenceMatcher(None, original, corrected)
    
    # Find all words in the corrected text
    corrected_words = list(re.finditer(r'\S+', corrected))

    for tag, i1, i2, j1, j2 in s.get_opcodes():
        if tag == 'equal':
            continue

        # Find which words in the corrected text (j1 to j2) were affected
        for word_match in corrected_words:
            start_char, end_char = word_match.span()
            # Check for overlap between the changed segment (j1, j2) and the word span
            if max(j1, start_char) < min(j2, end_char):
                # Highlight the word in the output box
                output_box.tag_add("correction", f"1.0+{start_char}c", f"1.0+{end_char}c")


def add_to_data_log(metrics, processing_time, language, style, model):
    """Add current correction to data log"""
    global data_log
    entry = {
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'words': metrics['word_count'],
        'time': f"{processing_time:.2f}",
        'f05': f"{metrics['f05']:.3f}",
        'err': f"{metrics['err']:.2f}%",
        'gleu': f"{metrics['gleu']:.3f}",
        'grade_level': f"{metrics['grade_level']:.2f}",
        'language': language,
        'style': style,
        'model': model
    }
    data_log.append(entry)
    update_data_table()

def update_data_table():
    """Update the data table with all logged entries"""
    # Clear existing data
    for row in data_tree.get_children():
        data_tree.delete(row)
    
    # Add all entries in reverse order (newest first)
    for entry in reversed(data_log):
        data_tree.insert("", "end", values=(
            entry['words'],
            entry['time'],
            entry['f05'],
            entry['err'],
            entry['gleu'],
            entry['grade_level'],
            entry['language'],
            entry['style'],
            entry['model']
        ))

def clear_data_log():
    """Clear all data from the log"""
    global data_log
    data_log = []
    update_data_table()
    messagebox.showinfo("Data Log", "All data cleared from log.")

def export_data_log():
    """Export data log to CSV file"""
    if not data_log:
        messagebox.showwarning("Export", "No data to export.")
        return
    
    filename = filedialog.asksaveasfilename(
        defaultextension=".csv",
        filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
        title="Export Data Log"
    )
    
    if filename:
        try:
            with open(filename, 'w', newline='', encoding='utf-8') as file:
                writer = csv.DictWriter(file, fieldnames=LOG_FIELDNAMES)
                writer.writeheader()
                writer.writerows(data_log)
            messagebox.showinfo("Export Successful", f"Data exported to {filename}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export data: {str(e)}")

def correct_grammar_threaded():
    def task():
        global start_time
        start_time = time.time()
        text = input_box.get("1.0", tk.END).strip()
        
        if not text:
            messagebox.showwarning("Warning", "Please enter some text."); return
            
        root.after(0, lambda: status_label.config(text="Processing...")); root.after(0, progress.start)
        
        try:
            lang = language_var.get()
            styl = style_var.get()
            model_used = model_var.get().split(" (Not installed)")[0]
            
            corr = run_correction(text, lang, styl)
            processing_time = time.time() - start_time
            
            # --- FIX: Calculate metrics before adding to log and updating output ---
            metrics = calculate_metrics(text, corr, processing_time)
            
            add_to_data_log(metrics, processing_time, lang, styl, model_used)
            root.after(0, lambda: update_output(corr, text, processing_time, metrics))
            
        except Exception as e:
            root.after(0, lambda: show_error(f"Correction failed: {e}", time.time() - start_time))
        finally:
            root.after(0, progress.stop)
            
    Thread(target=task, daemon=True).start()

# Updated update_output to correctly accept and use the 'metrics' argument
def update_output(c, o, t, m):
    output_box.delete("1.0",tk.END);output_box.insert(tk.END,c)
    highlight_corrections(o, c)
    
    # Use superscript for F⁰·⁵
    stats_label.config(
        text=f"Words:{m['word_count']} | Time:{t:.2f}s | F⁰·⁵:{m['f05']:.3f} | ERR:{m['err']:.2f}% | GLEU:{m['gleu']:.3f} | Grade Level:{m['grade_level']:.2f} | Score:{m['composite_score']:.1f}%"
    )
    status_label.config(text="Done.")

def show_error(m,t=0):messagebox.showerror("Error",m)

def clear_text():
    input_box.delete("1.0",tk.END);output_box.delete("1.0",tk.END)
    output_box.tag_delete("correction")
    # Updated initial stats_label to include Score
    stats_label.config(text="Words:0 | Time:0.00s | F⁰·⁵:0.000 | ERR:0.00% | GLEU:0.000 | Grade Level:0.00 ")

def copy_output():
    t=output_box.get("1.0",tk.END).strip()
    if not t:messagebox.showwarning("Warning","No text to copy!");return
    pyperclip.copy(t);messagebox.showinfo("Copied","Copied!")

def save_output():
    t=output_box.get("1.0",tk.END).strip()
    if not t:return messagebox.showwarning("Warning","No text to save!")
    f=filedialog.asksaveasfilename(defaultextension=".txt",filetypes=[("Text","*.txt"),("Word","*.docx")])
    if not f:return
    try:
        if f.endswith('.txt'):open(f,'w',encoding='utf-8').write(t)
        else:doc=Document();doc.add_paragraph(t);doc.save(f)
        messagebox.showinfo("Saved","File saved.")
    except Exception as e:messagebox.showerror("Error",str(e))

def swap_text():
    i=input_box.get("1.0",tk.END).strip();o=output_box.get("1.0",tk.END).strip()
    input_box.delete("1.0",tk.END);input_box.insert(tk.END,o)
    output_box.delete("1.0",tk.END);output_box.insert(tk.END,i)
    output_box.tag_delete("correction")

def load_file():
    f=filedialog.askopenfilename(filetypes=[("Text","*.txt"),("Word","*.docx"),("All","*.*")])
    if not f:return
    try:
        if f.endswith('.docx'):
            doc=Document(f);t='\n'.join(p.text for p in doc.paragraphs)
        else:t=open(f,'r',encoding='utf-8').read()
        input_box.delete("1.0",tk.END);input_box.insert(tk.END,t)
    except Exception as e:messagebox.showerror("Error",str(e))

# Create main window
root=tk.Tk();root.title("UALR Grammar Checker with Data Logger");root.geometry("1400x750")

# Create notebook for tabs
notebook = ttk.Notebook(root)
notebook.pack(fill='both', expand=True, padx=10, pady=10)

# Tab 1: Grammar Checker
tab1 = ttk.Frame(notebook)
notebook.add(tab1, text='Grammar Checker')

# Tab 2: Data Logger
tab2 = ttk.Frame(notebook)
notebook.add(tab2, text='Data Logger')

# Move all existing GUI elements to tab1
top_frame=tk.Frame(tab1);top_frame.pack(fill=tk.X)
controls_frame=tk.Frame(tab1);controls_frame.pack(fill=tk.X)
text_frame=tk.Frame(tab1);text_frame.pack(fill=tk.BOTH,expand=True)
bottom_frame=tk.Frame(tab1);bottom_frame.pack(fill=tk.X)

model_var=tk.StringVar(value=AVAILABLE_MODELS[0])
tk.Label(controls_frame,text="AI Model:").pack(side=tk.LEFT,padx=5)
model_menu=tk.OptionMenu(controls_frame,model_var,*AVAILABLE_MODELS)
model_menu.pack(side=tk.LEFT,padx=5)

tk.Button(controls_frame,text="🔄 Refresh Models",command=refresh_model_list).pack(side=tk.LEFT,padx=5)

def force_models_installed():
    model_var.set(AVAILABLE_MODELS[0]);messagebox.showinfo("Manual Override","Model set.")

tk.Button(controls_frame,text="🔧 Manual Override",command=force_models_installed).pack(side=tk.LEFT,padx=5)

language_var=tk.StringVar(value="English");style_var=tk.StringVar(value="Default")
tk.Label(controls_frame,text="Language:").pack(side=tk.LEFT,padx=5)
tk.OptionMenu(controls_frame,language_var,"English","Spanish","French","German","Italian","Portuguese","Dutch","Russian","Japanese","Chinese","Korean","Arabic","Hindi","Bengali").pack(side=tk.LEFT)
tk.Label(controls_frame,text="Style:").pack(side=tk.LEFT,padx=5)
tk.OptionMenu(controls_frame,style_var,"Default","Formal","Academic","Casual","Professional").pack(side=tk.LEFT)

input_label=tk.Label(text_frame,text="Input Text:");input_label.grid(row=0,column=0,sticky="w")
input_box=scrolledtext.ScrolledText(text_frame,width=60,height=20);input_box.grid(row=1,column=0,padx=5,pady=5,sticky="nsew")
output_label=tk.Label(text_frame,text="AI Corrected Text:");output_label.grid(row=0,column=1,sticky="w")
output_box=scrolledtext.ScrolledText(text_frame,width=60,height=20);output_box.grid(row=1,column=1,padx=5,pady=5,sticky="nsew")
text_frame.columnconfigure(0,weight=1);text_frame.columnconfigure(1,weight=1)

# Buttons without TTS
tk.Button(top_frame,text="Correct Grammar",command=correct_grammar_threaded).pack(side=tk.LEFT,padx=5)
tk.Button(top_frame,text="Clear",command=clear_text).pack(side=tk.LEFT,padx=5)
tk.Button(top_frame,text="Swap Text",command=swap_text).pack(side=tk.LEFT,padx=5)
tk.Button(top_frame,text="Copy",command=copy_output).pack(side=tk.LEFT,padx=5)
tk.Button(top_frame,text="Save",command=save_output).pack(side=tk.LEFT,padx=5)
tk.Button(top_frame,text="Load File",command=load_file).pack(side=tk.LEFT,padx=5)

progress=ttk.Progressbar(bottom_frame,mode='indeterminate');progress.pack(fill=tk.X,padx=5,pady=2)
status_label=tk.Label(bottom_frame,text="Ready. Click 'Refresh Models' if needed.",fg="blue");status_label.pack(anchor="w")

# Use superscript for F⁰·⁵ in initial stats, added Score
stats_label=tk.Label(bottom_frame,text="Words:0 | Time:0.00s | F⁰·⁵:0.000 | ERR:0.00% | GLEU:0.000 | Grade Level:0.00 ",fg="darkgreen",font=("Cambria",10,"bold"));stats_label.pack(anchor="w")

try:
    # Use a dummy image if ualr_logo1.png is not found to prevent full code failure
    # Create a small blank image instead of relying on a file
    img_size = (300, 100)
    img_array = np.zeros(img_size[::-1] + (3,), dtype=np.uint8) + 240 # Light gray background
    img = Image.fromarray(img_array, 'RGB')
    photo=ImageTk.PhotoImage(img)
    lbl=tk.Label(bottom_frame,image=photo);lbl.image=photo;lbl.pack(side=tk.RIGHT,padx=10,anchor="e")
except Exception as e:
    print("Logo not loaded, using blank image or skipping:", e)
    pass # Continue execution even if image fails

# Data Logger Tab (Tab2) setup
data_log_frame = tk.Frame(tab2)
data_log_frame.pack(fill='both', expand=True, padx=10, pady=10)

# Data log controls
data_controls = tk.Frame(data_log_frame)
data_controls.pack(fill='x', pady=(0, 10))

tk.Button(data_controls, text="Clear Data", command=clear_data_log).pack(side=tk.LEFT, padx=5)
tk.Button(data_controls, text="Export to CSV", command=export_data_log).pack(side=tk.LEFT, padx=5)

# Data table
columns = ('Words', 'Time', 'F⁰·⁵', 'ERR', 'GLEU', 'Grade Level','Language', 'Style', 'Model')
data_tree = ttk.Treeview(data_log_frame, columns=columns, show='headings', height=15)

# Define headings
data_tree.heading('Words', text='No of Words')
data_tree.heading('Time', text='Time (s)')
data_tree.heading('F⁰·⁵', text='F⁰·⁵')
data_tree.heading('ERR', text='ERR (%)')
data_tree.heading('GLEU', text='GLEU')
data_tree.heading('Grade Level', text='Grade Level')
data_tree.heading('Language', text='Language')
data_tree.heading('Style', text='Writing Style')
data_tree.heading('Model', text='LLM Model')

# Set column widths
data_tree.column('Words', width=80)
data_tree.column('Time', width=80)
data_tree.column('F⁰·⁵', width=80)
data_tree.column('ERR', width=80)
data_tree.column('GLEU', width=80)
data_tree.column('Grade Level', width=100)
data_tree.column('Language', width=100)
data_tree.column('Style', width=100)
data_tree.column('Model', width=120)

# Add scrollbar to data table
data_scrollbar = ttk.Scrollbar(data_log_frame, orient=tk.VERTICAL, command=data_tree.yview)
data_tree.configure(yscrollcommand=data_scrollbar.set)
data_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
data_tree.pack(side=tk.LEFT, fill='both', expand=True)

# Status label for data tab
data_status = tk.Label(
    data_log_frame, 
    text="Data log will automatically record each grammar correction.", 
    fg="blue"
)
data_status.pack(anchor='w', pady=(10, 0))


# --- Metric Definitions Section ---
metrics_frame = tk.LabelFrame(data_log_frame, text="Metric Definitions", padx=10, pady=10)
metrics_frame.pack(fill='x', pady=(10, 5))

# Text widget for formatting, color, and alignment
metrics_text_widget = tk.Text(
    metrics_frame,
    wrap="word",
    height=40, # Reduced height for better fit
    font=("Cambria", 10),
    relief="flat",
    bg=root.cget("bg"),
    padx=12,
    pady=8
)

# Define text tags for formatting
metrics_text_widget.tag_configure("title", foreground="#003366", font=("Cambria", 10, "bold"))
metrics_text_widget.tag_configure("formula", foreground="#444444", font=("Consolas", 10, "italic"))
metrics_text_widget.tag_configure("body", justify="left", lmargin1=20, lmargin2=35)

# --- Metric Descriptions with Equations ---
# --- F0.5 Metric (with clickable citation) ---
metrics_text_widget.insert("end", "• F⁰·⁵ Score (F0.5) ", "title")

# Insert clickable citation
start_index = metrics_text_widget.index("end")
metrics_text_widget.insert("end", "[Dahlmeier and Ng, 2012]\n", "f05_link")

# Style and hyperlink behavior
metrics_text_widget.tag_config("f05_link", foreground="blue", underline=True)
def open_f05_link(event):
    webbrowser.open_new("https://aclanthology.org/N12-1067.pdf")
metrics_text_widget.tag_bind("f05_link", "<Button-1>", open_f05_link)

# Equation and explanation
metrics_text_widget.insert(
    "end",
    "   Equation: F₀․₅ = (1 + 0.5²) × (P × R) / ((0.5² × P) + R)\n",
    "formula"
)
metrics_text_widget.insert(
    "end",
    "A precision-weighted F-score that gives twice as much importance to precision (P) as recall (R).\n"
    "It measures how accurately and completely the corrected text matches the intended meaning (P) and how many errors were corrected (R).\n\n",
    "body"
)

# --- GLEU metric with clickable citation ---
metrics_text_widget.insert("end", "• GLEU – Google’s Grammar/Translation Edit Metric ", "title")

metrics_text_widget.insert("end", "[Napoles et al., 2015]\n", "gleu_link")
metrics_text_widget.tag_config("gleu_link", foreground="blue", underline=True)
def open_gleu_link(event):
    webbrowser.open_new("https://doi.org/10.48550/arXiv.1605.02592")
metrics_text_widget.tag_bind("gleu_link", "<Button-1>", open_gleu_link)

metrics_text_widget.insert(
    "end",
    "   Equation: GLEU = (Matching n-grams) / (Total n-grams)\n",
    "formula"
)
metrics_text_widget.insert(
    "end",
    "Compares n-gram overlap between the corrected and reference text. Higher GLEU indicates better fluency and similarity to the original style.\n\n",
    "body"
)

# --- ERR metric with clickable citation ---
metrics_text_widget.insert("end", "• ERR (%) – Error Rate Reduction ", "title")

metrics_text_widget.insert("end", "[Jin Wang 2025]\n", "err_link")
metrics_text_widget.tag_config("err_link", foreground="blue", underline=True)
def open_err_link(event):
    webbrowser.open_new("https://doi.org/10.1016/j.aej.2025.08.005")
metrics_text_widget.tag_bind("err_link", "<Button-1>", open_err_link)

metrics_text_widget.insert(
    "end",
    "   Equation: ERR = (Edits / Total_Words) × 100\n",
    "formula"
)
metrics_text_widget.insert(
    "end",
    "Represents the percentage of edits relative to the total number of words. A lower ERR indicates fewer unnecessary or incorrect changes.\n\n",
    "body"
)

# --- Grade Score metric with clickable citation (Updated description to match Flesch-Kincaid) ---
metrics_text_widget.insert("end", "• Flesch-Kincaid Grade Level ", "title")

metrics_text_widget.insert("end", "[U.S Navy 1975]\n", "grade_link")
metrics_text_widget.tag_config("grade_link", foreground="blue", underline=True)
def open_grade_link(event):
    webbrowser.open_new("https://en.wikipedia.org/wiki/Flesch%E2%80%93Kincaid_readability_tests")
metrics_text_widget.tag_bind("grade_link", "<Button-1>", open_grade_link)
metrics_text_widget.insert(
    "end",
    "   Equation: Grade Level = 0.39 × (Words/Sentences) + 11.8 × (Syllables/Words) - 15.59\n",
    "formula"
)

metrics_text_widget.insert(
    "end",
    "A readability test that estimates the U.S. school grade level needed to understand the text.\n"
    "A score of 8.0, for example, means a text is readable by an average 8th grader.\n\n",
    "body"
)

metrics_text_widget.configure(state="disabled")
metrics_text_widget.pack(fill="x", expand=True)


# Initial model refresh
refresh_model_list()
try:
    img=Image.open("ualr_logo1.png");img=img.resize((300,100));photo=ImageTk.PhotoImage(img)
    lbl=tk.Label(bottom_frame,image=photo);lbl.image=photo;lbl.pack(side=tk.RIGHT,padx=20,anchor="e")
except Exception as e:print("Logo not loaded",e)
root.mainloop()
